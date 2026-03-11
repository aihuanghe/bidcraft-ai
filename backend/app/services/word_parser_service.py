"""Word文档解析服务"""
import os
import re
import json
import asyncio
import aiofiles
from typing import Optional, List, Dict, Any, AsyncGenerator, Tuple
from datetime import datetime
from pathlib import Path

import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from ..config import settings


class WordParseService:
    """Word文档解析服务"""
    
    # 大文件阈值（50MB）
    LARGE_FILE_THRESHOLD = 50 * 1024 * 1024
    
    # 支持的Word格式
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']
    
    @staticmethod
    def is_supported(file_path: str) -> bool:
        """检查是否支持该文件格式"""
        ext = os.path.splitext(file_path)[1].lower()
        return ext in WordParseService.SUPPORTED_EXTENSIONS
    
    @staticmethod
    async def parse_word_stream(
        file_path: str,
        progress_callback: Optional[Callable[[int, str], None]] = None
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """
        流式解析Word文档，按段落yield内容
        Yields: {
            "type": "paragraph" | "heading" | "table" | "image",
            "level": 1-6,  # 标题级别
            "content": str,
            "style": dict,  # 样式信息
            "index": int
        }
        """
        try:
            # 获取文件大小
            file_size = os.path.getsize(file_path)
            is_large_file = file_size > WordParseService.LARGE_FILE_THRESHOLD
            
            if is_large_file:
                print(f"大文件检测: {file_size / 1024 / 1024:.2f}MB，采用分段解析模式")
            
            # 提取格式规则
            format_rules = await WordParseService.extract_format_rules(file_path)
            if progress_callback:
                progress_callback(5, "提取格式规则完成")
            
            yield {
                "type": "format_rules",
                "content": format_rules
            }
            
            # 解析文档结构
            doc = docx.Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            total_tables = len(doc.tables)
            total_items = total_paragraphs + total_tables
            
            processed = 0
            heading_tree = []  # 标题树结构
            current_section = None
            
            # 处理段落
            for i, para in enumerate(doc.paragraphs):
                # 检测是否为标题
                heading_level = WordParseService._detect_heading_level(para)
                
                if heading_level > 0:
                    # 标题段落
                    heading_info = {
                        "level": heading_level,
                        "text": para.text.strip(),
                        "index": i
                    }
                    heading_tree.append(heading_info)
                    current_section = heading_info
                    
                    yield {
                        "type": "heading",
                        "level": heading_level,
                        "content": para.text.strip(),
                        "style": WordParseService._extract_paragraph_style(para),
                        "index": i
                    }
                else:
                    # 普通段落
                    text = para.text.strip()
                    if text:  # 跳过空段落
                        yield {
                            "type": "paragraph",
                            "content": text,
                            "style": WordParseService._extract_paragraph_style(para),
                            "index": i,
                            "section": current_section.get("text") if current_section else None
                        }
                
                processed += 1
                
                # 进度更新（20%-70%）
                if progress_callback and total_items > 0:
                    progress = 20 + int(50 * processed / total_items)
                    progress_callback(progress, f"已处理 {processed}/{total_items}")
                
                # 大文件分段处理，释放内存
                if is_large_file and i > 0 and i % 500 == 0:
                    await asyncio.sleep(0.1)  # 让出控制权
            
            # 处理表格
            for i, table in enumerate(doc.tables):
                table_data = WordParseService._extract_table_data(table)
                yield {
                    "type": "table",
                    "content": table_data,
                    "rows": len(table.rows),
                    "cols": len(table.columns) if table.rows else 0,
                    "index": total_paragraphs + i
                }
                
                processed += 1
                if progress_callback and total_items > 0:
                    progress = 70 + int(25 * processed / total_items)
                    progress_callback(progress, f"已处理 {processed}/{total_items}")
            
            # 提取嵌入图片（OCR）
            if progress_callback:
                progress_callback(95, "检测嵌入图片...")
            
            images = await WordParseService.extract_embedded_images(file_path)
            if images:
                yield {
                    "type": "embedded_images",
                    "count": len(images),
                    "content": images
                }
            
            # 完成
            yield {
                "type": "complete",
                "content": {
                    "total_paragraphs": total_paragraphs,
                    "total_tables": total_tables,
                    "total_images": len(images),
                    "heading_tree": heading_tree
                }
            }
            
            if progress_callback:
                progress_callback(100, "解析完成")
                
        except Exception as e:
            yield {
                "type": "error",
                "content": str(e)
            }
    
    @staticmethod
    def _detect_heading_level(para: docx.text.paragraph.Paragraph) -> int:
        """检测段落是否为标题及其级别"""
        # 方法1: 检查样式名称
        style_name = para.style.name.lower() if para.style else ""
        if "heading" in style_name or "标题" in style_name:
            # 提取级别数字
            match = re.search(r'(\d+)', style_name)
            if match:
                return int(match.group(1))
            return 1  # 默认一级标题
        
        # 方法2: 检查字体格式（粗体+大字号）
        if para.runs:
            first_run = para.runs[0]
            font_size = first_run.font.size
            is_bold = first_run.font.bold
            
            # 判断是否为标题格式
            if font_size and font_size Pt(14):  # 四号及以上
                if is_bold:
                    # 进一步分析位置和格式
                    return 1
                elif font_size >= Pt(16):  # 二号及以上
                    return 1
        
        return 0  # 普通段落
    
    @staticmethod
    def _extract_paragraph_style(para: docx.text.paragraph.Paragraph) -> Dict[str, Any]:
        """提取段落样式信息"""
        style = {
            "font_name": "宋体",
            "font_size": "小四",  # 默认
            "is_bold": False,
            "is_italic": False,
            "alignment": "left",
            "line_spacing": 1.5,
            "first_line_indent": 0,
        }
        
        if not para.runs:
            return style
        
        # 提取字体信息（取第一个run的格式）
        first_run = para.runs[0]
        
        if first_run.font.name:
            style["font_name"] = first_run.font.name
        
        if first_run.font.size:
            pt_size = first_run.font.size.pt
            style["font_size"] = WordParseService._pt_to_size_name(pt_size)
        
        style["is_bold"] = first_run.font.bold or False
        style["is_italic"] = first_run.font.italic or False
        
        # 对齐方式
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
        }
        style["alignment"] = alignment_map.get(para.alignment, "left")
        
        # 行距
        if para.paragraph_format.line_spacing_rule:
            style["line_spacing"] = para.paragraph_format.line_spacing
        
        # 首行缩进
        if para.paragraph_format.first_line_indent:
            style["first_line_indent"] = para.paragraph_format.first_line_indent.pt
        
        return style
    
    @staticmethod
    def _pt_to_size_name(pt_size: float) -> str:
        """将字号转换为中文名称"""
        size_map = {
            42: "初号", 36: "一号", 31.5: "二号", 28: "三号",
            24: "四号", 21: "五号", 18: "小六", 16: "六号",
            14: "小四", 12: "五号", 10.5: "小五", 9: "六号",
            7.5: "七号", 6.5: "八号", 5.5: "九号"
        }
        
        for size, name in size_map.items():
            if abs(pt_size - size) < 2:
                return name
        return f"{pt_size}pt"
    
    @staticmethod
    def _extract_table_data(table: docx.table.Table) -> List[List[str]]:
        """提取表格数据为二维数组"""
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_data.append(cell_text)
            data.append(row_data)
        return data
    
    @staticmethod
    async def extract_format_rules(file_path: str) -> Dict[str, Any]:
        """从Word文档中提取格式规则"""
        rules = {
            "fonts": {
                "body": {"name": "宋体", "size": "小四"},
                "heading1": {"name": "黑体", "size": "三号", "bold": True},
                "heading2": {"name": "黑体", "size": "四号", "bold": True},
                "heading3": {"name": "黑体", "size": "小四", "bold": True},
            },
            "paragraph": {
                "line_spacing": 1.5,
                "first_line_indent": 2,
            },
            "page": {
                "top_margin": 2.54,
                "bottom_margin": 2.54,
                "left_margin": 3.17,
                "right_margin": 3.17,
                "paper_size": "A4",
            }
        }
        
        try:
            doc = docx.Document(file_path)
            
            # 从文档样式中提取格式
            if doc.styles:
                # 正文样式
                normal_style = doc.styles.get('Normal')
                if normal_style and normal_style._element:
                    # 提取字体
                    pass
                
                # 标题样式
                for level in range(1, 4):
                    heading_style = doc.styles.get(f'Heading {level}')
                    if heading_style:
                        # 提取标题格式
                        pass
            
            # 从节设置中提取页面格式
            if doc.sections:
                section = doc.sections[0]
                rules["page"]["top_margin"] = section.top_margin.cm if section.top_margin else 2.54
                rules["page"]["bottom_margin"] = section.bottom_margin.cm if section.bottom_margin else 2.54
                rules["page"]["left_margin"] = section.left_margin.cm if section.left_margin else 3.17
                rules["page"]["right_margin"] = section.right_margin.cm if section.right_margin else 3.17
                
                # 纸张大小
                if section.page_width:
                    width_cm = section.page_width.cm
                    if width_cm:
                        rules["page"]["paper_size"] = "A4" if abs(width_cm - 21) < 2 else "其他"
            
        except Exception as e:
            print(f"提取格式规则失败: {e}")
        
        return rules
    
    @staticmethod
    async def extract_embedded_images(file_path: str) -> List[Dict[str, Any]]:
        """提取Word文档中的嵌入图片"""
        images = []
        
        try:
            doc = docx.Document(file_path)
            
            # 遍历所有内联图形
            for para in doc.paragraphs:
                for run in para.runs:
                    if hasattr(run, '_element'):
                        # 查找inlineShape元素
                        for elem in run._element.iter():
                            if elem.tag.endswith('}inlineShape') or elem.tag.endswith('}pict'):
                                # 找到图片
                                images.append({
                                    "type": "embedded",
                                    "paragraph_index": doc.paragraphs.index(para),
                                    "text_before": para.text[:50]  # 图片前的文字
                                })
            
            # 遍历所有表格中的图片
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if hasattr(run, '_element'):
                                    for elem in run._element.iter():
                                        if elem.tag.endswith('}inlineShape'):
                                            images.append({
                                                "type": "table_image",
                                                "text": para.text[:30]
                                            })
            
        except Exception as e:
            print(f"提取嵌入图片失败: {e}")
        
        return images
    
    @staticmethod
    async def parse_word_full(
        file_path: str,
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """
        完整解析Word文档，返回结构化数据
        """
        result = {
            "file_path": file_path,
            "file_size": os.path.getsize(file_path),
            "parsed_at": datetime.utcnow().isoformat(),
            "format_rules": None,
            "structure": {
                "headings": [],
                "paragraphs": [],
                "tables": [],
            },
            "images": [],
            "errors": []
        }
        
        try:
            # 收集所有内容
            headings = []
            paragraphs = []
            tables = []
            
            async for item in WordParseService.parse_word_stream(file_path, progress_callback):
                item_type = item.get("type")
                
                if item_type == "format_rules":
                    result["format_rules"] = item.get("content")
                elif item_type == "heading":
                    headings.append({
                        "level": item.get("level"),
                        "text": item.get("content"),
                        "index": item.get("index")
                    })
                elif item_type == "paragraph":
                    paragraphs.append({
                        "text": item.get("content"),
                        "section": item.get("section"),
                        "index": item.get("index")
                    })
                elif item_type == "table":
                    tables.append({
                        "data": item.get("content"),
                        "rows": item.get("rows"),
                        "cols": item.get("cols"),
                        "index": item.get("index")
                    })
                elif item_type == "embedded_images":
                    result["images"] = item.get("content", [])
                elif item_type == "error":
                    result["errors"].append(item.get("content"))
            
            result["structure"] = {
                "headings": headings,
                "paragraphs": paragraphs[:1000],  # 限制返回数量
                "tables": tables,
            }
            
        except Exception as e:
            result["errors"].append(str(e))
        
        return result


# 类型别名
from typing import Callable, Any, Dict

word_parser_service = WordParseService()
