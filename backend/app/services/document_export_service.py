"""文档导出服务 - 基于模板的Word文档生成"""
import os
import json
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from sqlalchemy.orm import Session

from ..models.models import BidProject, ExtractedTemplate, DocumentOutline, DocumentContent


class DocumentExportService:
    """基于模板的Word文档导出服务"""
    
    DEFAULT_FONTS = {
        "title": {"name": "黑体", "size": 22, "bold": True},
        "heading1": {"name": "黑体", "size": 16, "bold": True},
        "heading2": {"name": "黑体", "size": 14, "bold": True},
        "heading3": {"name": "楷体", "size": 12, "bold": True},
        "body": {"name": "宋体", "size": 11, "bold": False},
    }
    
    def __init__(self, db: Session):
        self.db = db
    
    def export_bid_document(
        self,
        project: BidProject,
        outline: Dict[str, Any],
        contents: Dict[str, str],
        output_path: str = None
    ) -> str:
        """导出投标文档
        
        Args:
            project: 投标项目
            outline: 大纲结构
            contents: 章节内容字典 {chapter_id: content}
            output_path: 输出路径
        
        Returns:
            导出文件路径
        """
        doc = Document()
        
        self._apply_page_settings(doc)
        
        self._add_cover_page(doc, project)
        
        self._build_document_from_outline(doc, outline, contents)
        
        self._add_table_of_contents(doc, outline)
        
        if output_path is None:
            output_dir = os.path.join(os.getcwd(), "exports")
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, f"{project.name}_投标文件.docx")
        
        doc.save(output_path)
        
        return output_path
    
    def _apply_page_settings(self, doc: Document):
        """应用页面设置"""
        section = doc.sections[0]
        
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    def _add_cover_page(self, doc: Document, project: BidProject):
        """添加封面页"""
        section = doc.add_section()
        section.top_margin = Inches(2.5)
        
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(project.name or "投标文件")
        title_run.font.name = "黑体"
        title_run.font.size = Pt(22)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        company_para = doc.add_paragraph()
        company_run = company_para.add_run(f"投标人：{{{{company_name}}}}")
        company_run.font.name = "宋体"
        company_run.font.size = Pt(16)
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"日期：{{{{bid_date}}}}")
        date_run.font.name = "宋体"
        date_run.font.size = Pt(14)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        section = doc.sections[-1]
        section.top_margin = Inches(1.0)
    
    def _build_document_from_outline(
        self,
        doc: Document,
        outline: Dict[str, Any],
        contents: Dict[str, str]
    ):
        """根据大纲构建文档内容"""
        chapters = outline.get("chapters", [])
        
        for chapter in chapters:
            self._add_chapter_content(doc, chapter, contents)
    
    def _add_chapter_content(
        self,
        doc: Document,
        chapter: Dict[str, Any],
        contents: Dict[str, str],
        parent_level: int = 0
    ):
        """添加章节内容"""
        level = chapter.get("level", 1)
        
        if level == 1:
            self._add_heading1(doc, chapter["title"])
        elif level == 2:
            self._add_heading2(doc, chapter["title"])
        else:
            self._add_heading3(doc, chapter["title"])
        
        template_snippet = chapter.get("template_snippet", "")
        if template_snippet:
            self._add_template_format(doc, template_snippet)
        
        chapter_id = chapter.get("id", "")
        if chapter_id in contents:
            content = contents[chapter_id]
            self._add_content_paragraph(doc, content)
        
        if chapter.get("has_table"):
            self._add_table_from_format(doc, chapter.get("table_format", {}))
        
        children = chapter.get("children", [])
        for child in children:
            self._add_chapter_content(doc, child, contents, level)
    
    def _add_heading1(self, doc: Document, text: str):
        """添加一级标题"""
        para = doc.add_paragraph(text)
        para.style = "Heading 1"
        for run in para.runs:
            run.font.name = "黑体"
            run.font.size = Pt(16)
            run.bold = True
    
    def _add_heading2(self, doc: Document, text: str):
        """添加二级标题"""
        para = doc.add_paragraph(text)
        para.style = "Heading 2"
        for run in para.runs:
            run.font.name = "黑体"
            run.font.size = Pt(14)
            run.bold = True
    
    def _add_heading3(self, doc: Document, text: str):
        """添加三级标题"""
        para = doc.add_paragraph(text)
        para.style = "Heading 3"
        for run in para.runs:
            run.font.name = "楷体"
            run.font.size = Pt(12)
            run.bold = True
    
    def _add_template_format(self, doc: Document, template: str):
        """添加模板格式说明"""
        para = doc.add_paragraph()
        run = para.add_run(f"【模板格式】{template[:200]}")
        run.font.name = "宋体"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_content_paragraph(self, doc: Document, content: str):
        """添加内容段落"""
        paragraphs = content.split("\n\n")
        
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                for run in para.runs:
                    run.font.name = "宋体"
                    run.font.size = Pt(11)
    
    def _add_table_from_format(self, doc: Document, table_format: Dict[str, Any]):
        """根据表格格式添加表格"""
        columns = table_format.get("columns", [])
        if not columns:
            return
        
        table = doc.add_table(rows=2, cols=len(columns))
        table.style = "Table Grid"
        
        header_row = table.rows[0]
        for idx, col_name in enumerate(columns):
            cell = header_row.cells[idx]
            cell.text = col_name
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "黑体"
                    run.font.size = Pt(10)
                    run.bold = True
        
        data_row = table.rows[1]
        for idx in range(len(columns)):
            cell = data_row.cells[idx]
            cell.text = ""
    
    def _add_table_of_contents(self, doc: Document, outline: Dict[str, Any]):
        """添加目录"""
        doc.add_page_break()
        
        toc_title = doc.add_paragraph("目 录")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in toc_title.runs:
            run.font.name = "黑体"
            run.font.size = Pt(16)
            run.bold = True
        
        doc.add_paragraph()
        
        chapters = outline.get("chapters", [])
        self._add_toc_entries(doc, chapters, 0)
    
    def _add_toc_entries(self, doc: Document, chapters: List[Dict[str, Any]], level: int):
        """添加目录条目"""
        for chapter in chapters:
            title = chapter.get("title", "")
            chapter_id = chapter.get("chapter_id", "")
            
            toc_para = doc.add_paragraph()
            
            if level == 0:
                toc_run = toc_para.add_run(f"{chapter_id} {title}")
                toc_run.font.name = "黑体"
                toc_run.font.size = Pt(14)
            else:
                indent = "  " * level
                toc_run = toc_para.add_run(f"{indent}{chapter_id} {title}")
                toc_run.font.name = "宋体"
                toc_run.font.size = Pt(12)
            
            children = chapter.get("children", [])
            if children:
                self._add_toc_entries(doc, children, level + 1)
    
    def validate_document_format(
        self,
        doc: Document,
        template: ExtractedTemplate = None
    ) -> Dict[str, Any]:
        """验证文档格式合规性"""
        validation_result = {
            "passed": True,
            "issues": [],
            "warnings": []
        }
        
        style_rules = template.style_rules if template else {}
        
        page_rules = style_rules.get("page", {})
        if page_rules:
            section = doc.sections[0]
            
            if page_rules.get("top_margin"):
                actual_top = section.top_margin.cm
                expected_top = page_rules["top_margin"]
                if abs(actual_top - expected_top) > 0.1:
                    validation_result["issues"].append(
                        f"上边距不匹配: 期望{expected_top}cm, 实际{actual_top:.2f}cm"
                    )
        
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                for run in para.runs:
                    if run.font.name and run.font.name not in ["黑体", "楷体", "宋体"]:
                        validation_result["warnings"].append(
                            f"标题使用了非标准字体: {run.font.name}"
                        )
        
        if validation_result["issues"]:
            validation_result["passed"] = False
        
        return validation_result


def get_export_service(db: Session) -> DocumentExportService:
    """获取文档导出服务实例"""
    return DocumentExportService(db)