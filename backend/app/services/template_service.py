"""模板识别与提取服务"""
import re
from typing import Dict, List, Optional, Any
from datetime import datetime
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from sqlalchemy.orm import Session

from ..models.models import ExtractedTemplate, TenderDocument
from ..models.database import get_db


class TemplateExtractionService:
    """招标文件模板提取服务"""
    
    TEMPLATE_KEYWORDS = [
        "投标文件格式", "投标文件编制要求", "投标书格式", 
        "投标文件组成", "投标文件结构", "投标文件内容"
    ]
    
    TEMPLATE_SECTION_PATTERNS = [
        r"第[一二三四五六七八九十\d]+章\s*投标文件格式",
        r"第[一二三四五六七八九十\d]+章\s*投标文件编制要求",
        r"第[一二三四五六七八九十\d]+章\s*投标书格式",
        r"第[一二三四五六七八九十\d]+章\s*投标文件组成",
        r"第[一二三四五六七八九十\d]+[章节]\s*投标文件格式",
        r"第[一二三四五六七八九十\d]+[章节]\s*投标文件编制",
    ]
    
    FORMAT_SUBSECTION_KEYWORDS = {
        "letter": ["投标函", "致函", "投标书"],
        "business": ["商务部分", "资格审查", "业绩表", "财务表", "资质"],
        "technical": ["技术部分", "技术方案", "实施计划", "人员配置", "技术响应"],
        "price": ["报价部分", "报价表", "分项报价", "价格"],
    }
    
    def __init__(self, db: Session):
        self.db = db
    
    def detect_template_chapter(self, content: str) -> Optional[Dict[str, Any]]:
        """检测招标文件是否包含投标文件格式章节"""
        for keyword in self.TEMPLATE_KEYWORDS:
            if keyword in content:
                match = re.search(r"第([一二三四五六七八九十\d]+)[章节]", content[:content.find(keyword) + 20])
                if match:
                    return {
                        "found": True,
                        "chapter": match.group(0),
                        "chapter_num": match.group(1),
                        "keyword": keyword
                    }
        
        for pattern in self.TEMPLATE_SECTION_PATTERNS:
            match = re.search(pattern, content)
            if match:
                return {
                    "found": True,
                    "chapter": match.group(0),
                    "chapter_num": match.group(1) if match.lastgroup else None,
                    "keyword": match.group(0)
                }
        
        return {"found": False}
    
    def extract_template_from_docx(self, file_path: str) -> Optional[Dict[str, Any]]:
        """从Word文档中提取模板内容"""
        try:
            doc = Document(file_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
            template_info = self.detect_template_chapter(full_text)
            if not template_info or not template_info.get("found"):
                return None
            
            template_data = {
                "detected": True,
                "chapter_info": template_info,
                "sections": [],
                "style_rules": self._extract_style_rules(doc),
                "original_snippets": []
            }
            
            in_template_section = False
            current_section = None
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                if any(kw in text for kw in self.TEMPLATE_KEYWORDS):
                    in_template_section = True
                    continue
                
                if in_template_section:
                    section_info = self._classify_section(text)
                    if section_info:
                        if current_section:
                            template_data["sections"].append(current_section)
                        current_section = section_info
                    elif current_section and text:
                        current_section["content"] = current_section.get("content", "") + "\n" + text
                    
                    if self._is_likely_end_of_template(text, template_data["sections"]):
                        if current_section:
                            template_data["sections"].append(current_section)
                        break
            
            if current_section:
                template_data["sections"].append(current_section)
            
            template_data["confidence"] = self._calculate_confidence(template_data["sections"])
            
            return template_data
            
        except Exception as e:
            print(f"提取模板失败: {e}")
            return None
    
    def _extract_style_rules(self, doc: Document) -> Dict[str, Any]:
        """提取文档样式规则"""
        rules = {
            "fonts": {
                "title": None,
                "body": None
            },
            "paragraph": {
                "line_spacing": None,
                "first_line_indent": None
            },
            "page": {
                "top_margin": None,
                "bottom_margin": None,
                "left_margin": None,
                "right_margin": None
            }
        }
        
        if doc.styles:
            try:
                normal_style = doc.styles['Normal']
                if normal_style and normal_style.font:
                    rules["fonts"]["body"] = {
                        "name": normal_style.font.name,
                        "size": normal_style.font.size.pt if normal_style.font.size else None
                    }
            except:
                pass
        
        if doc.sections:
            section = doc.sections[0]
            rules["page"] = {
                "top_margin": section.top_margin.cm if section.top_margin else None,
                "bottom_margin": section.bottom_margin.cm if section.bottom_margin else None,
                "left_margin": section.left_margin.cm if section.left_margin else None,
                "right_margin": section.right_margin.cm if section.right_margin else None
            }
        
        return rules
    
    def _classify_section(self, text: str) -> Optional[Dict[str, Any]]:
        """分类模板章节"""
        text_lower = text.lower()
        
        for category, keywords in self.FORMAT_SUBSECTION_KEYWORDS.items():
            for kw in keywords:
                if kw in text:
                    return {
                        "type": category,
                        "title": text[:100],
                        "content": "",
                        "has_table": self._check_section_for_table(text),
                        "requirements": []
                    }
        
        if re.match(r"^[一（一）\d\.（\d）]+", text):
            return {
                "type": "other",
                "title": text[:100],
                "content": "",
                "has_table": False,
                "requirements": []
            }
        
        return None
    
    def _check_section_for_table(self, text: str) -> bool:
        """检查章节是否包含表格要求"""
        table_keywords = ["表", "表格", "格式", "列", "填写"]
        return any(kw in text for kw in table_keywords)
    
    def _is_likely_end_of_template(self, text: str, sections: List) -> bool:
        """判断是否到达模板区域结束"""
        if not text:
            return True
        
        end_indicators = ["附件", "附录", "注：本", "说明："]
        if any(indicator in text for indicator in end_indicators):
            return True
        
        if len(sections) > 10:
            return True
        
        return False
    
    def _calculate_confidence(self, sections: List) -> float:
        """计算模板提取置信度"""
        if not sections:
            return 0.0
        
        has_letter = any(s.get("type") == "letter" for s in sections)
        has_business = any(s.get("type") == "business" for s in sections)
        has_technical = any(s.get("type") == "technical" for s in sections)
        has_price = any(s.get("type") == "price" for s in sections)
        
        score = 0.3
        if has_letter:
            score += 0.2
        if has_business:
            score += 0.15
        if has_technical:
            score += 0.2
        if has_price:
            score += 0.15
        
        return min(score, 1.0)
    
    def save_extracted_template(
        self, 
        source_doc_id: int, 
        template_data: Dict[str, Any],
        industry: str = "general"
    ) -> ExtractedTemplate:
        """保存提取的模板到数据库"""
        template = ExtractedTemplate(
            source_doc_id=source_doc_id,
            template_type="extracted",
            name=f"从招标文件提取的模板 - {datetime.now().strftime('%Y%m%d')}",
            description=f"从招标文件提取的投标文件格式模板",
            industry=industry,
            structure_json={
                "sections": template_data.get("sections", []),
                "chapter_info": template_data.get("chapter_info", {})
            },
            style_rules=template_data.get("style_rules", {}),
            original_snippets=template_data.get("original_snippets", []),
            confidence_score=template_data.get("confidence", 0.5),
            is_active=True
        )
        
        self.db.add(template)
        self.db.commit()
        self.db.refresh(template)
        
        return template
    
    def update_document_template_flag(self, doc_id: int, template_id: int, chapter: str):
        """更新招标文件的模板标记"""
        doc = self.db.query(TenderDocument).filter(TenderDocument.id == doc_id).first()
        if doc:
            doc.has_format_template = True
            doc.format_template_chapter = chapter
            doc.extracted_template_id = template_id
            self.db.commit()
    
    def get_extracted_template(self, template_id: int) -> Optional[ExtractedTemplate]:
        """获取提取的模板"""
        return self.db.query(ExtractedTemplate).filter(
            ExtractedTemplate.id == template_id,
            ExtractedTemplate.is_active == True
        ).first()
    
    def get_templates_by_industry(self, industry: str) -> List[ExtractedTemplate]:
        """根据行业获取模板"""
        return self.db.query(ExtractedTemplate).filter(
            ExtractedTemplate.industry == industry,
            ExtractedTemplate.is_active == True
        ).all()


class TemplateMatchingService:
    """模板匹配服务"""
    
    INDUSTRY_KEYWORDS = {
        "engineering": ["工程", "施工", "建设", "建筑", "装修", "市政", "公路", "桥梁"],
        "it": ["软件", "信息化", "系统", "网络", "集成", "IT", "智慧", "数据"],
        "medical": ["医疗", "医院", "医药", "设备", "器械", "保健"],
        "government": ["政府采购", "办公", "服务", "物业", "租赁"],
    }
    
    def __init__(self, db: Session):
        self.db = db
    
    def detect_industry(self, project_overview: str, technical_requirements: str) -> str:
        """检测项目行业类型"""
        text = (project_overview or "") + (technical_requirements or "")
        
        for industry, keywords in self.INDUSTRY_KEYWORDS.items():
            if any(kw in text for kw in keywords):
                return industry
        
        return "general"
    
    def calculate_match_score(
        self, 
        template: ExtractedTemplate,
        industry: str,
        technical_scores: Dict[str, Any]
    ) -> Dict[str, Any]:
        """计算模板匹配分数"""
        score = 0.0
        reasons = []
        
        if template.template_type == "extracted":
            score += 0.4
            reasons.append("从招标文件提取的模板，格式完全匹配")
        
        if template.industry == industry:
            score += 0.3
            reasons.append(f"行业匹配: {industry}")
        elif template.industry == "general":
            score += 0.1
            reasons.append("通用模板，适用范围广")
        
        structure = template.structure_json or {}
        sections = structure.get("sections", [])
        
        has_all_sections = all([
            any(s.get("type") == "letter" for s in sections),
            any(s.get("type") == "business" for s in sections),
            any(s.get("type") == "technical" for s in sections),
        ])
        
        if has_all_sections:
            score += 0.2
            reasons.append("模板包含完整的章节结构")
        
        if template.confidence_score:
            score += template.confidence_score * 0.1
            reasons.append(f"模板置信度: {template.confidence_score:.2f}")
        
        return {
            "template_id": template.id,
            "score": min(score, 1.0),
            "reasons": reasons,
            "recommendation": self._get_recommendation(score)
        }
    
    def _get_recommendation(self, score: float) -> str:
        """根据分数给出推荐建议"""
        if score >= 0.8:
            return "直接使用"
        elif score >= 0.5:
            return "建议确认"
        else:
            return "建议补充"


def get_template_service(db: Session) -> TemplateExtractionService:
    """获取模板提取服务实例"""
    return TemplateExtractionService(db)


def get_matching_service(db: Session) -> TemplateMatchingService:
    """获取模板匹配服务实例"""
    return TemplateMatchingService(db)